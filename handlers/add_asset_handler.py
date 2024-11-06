from docx import Document
import os

output_directory = 'путь для сохранения файлов "Акт"'
os.makedirs(output_directory, exist_ok=True)

def create_asset_issue_act(data, file_name):

    doc = Document("template.docx")
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if f"{{{key}}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", str(value))

    output_file = os.path.join(output_directory, f"{file_name}.docx")
    doc.save(output_file)
    return output_file

def add_asset(bot, conn, message):
    user_data = {}

    bot.send_message(message.chat.id, "Введите номер акта:")
    bot.register_next_step_handler(message, lambda msg: get_act_number(bot, conn, msg, user_data))

def get_act_number(bot, conn, message, user_data):
    user_data['act_number'] = message.text
    bot.send_message(message.chat.id, "Сколько единиц оборудования вы хотите выдать?")
    bot.register_next_step_handler(message, lambda msg: get_equipment_count(bot, conn, msg, user_data))

def get_equipment_count(bot, conn, message, user_data):
    user_data['equipment_count'] = int(message.text)
    bot.send_message(message.chat.id, "Введите наименование оборудования:")
    bot.register_next_step_handler(message, lambda msg: get_asset_name(bot, conn, msg, user_data))

def get_asset_name(bot, conn, message, user_data):
    user_data['asset_name'] = message.text
    bot.send_message(message.chat.id, "Введите должность сотрудника:")
    bot.register_next_step_handler(message, lambda msg: get_position(bot, conn, msg, user_data))

def get_position(bot, conn, message, user_data):
    user_data['position'] = message.text
    bot.send_message(message.chat.id, "Введите ФИО сотрудника:")
    bot.register_next_step_handler(message, lambda msg: get_full_name(bot, conn, msg, user_data))

def get_full_name(bot, conn, message, user_data):
    user_data['full_name'] = message.text
    bot.send_message(message.chat.id, "Введите дату (ДД-ММ-ГГГГ):")
    bot.register_next_step_handler(message, lambda msg: get_date(bot, conn, msg, user_data))

def get_date(bot, conn, message, user_data):
    user_data['date'] = message.text
    bot.send_message(message.chat.id, "Введите отдел:")
    bot.register_next_step_handler(message, lambda msg: get_department(bot, conn, msg, user_data))

def get_department(bot, conn, message, user_data):
    user_data['department'] = message.text
    bot.send_message(message.chat.id, "Введите дату выпуска оборудования (ДД-ММ-ГГГГ):")
    bot.register_next_step_handler(message, lambda msg: get_equipment_release_date(bot, conn, msg, user_data))

def get_equipment_release_date(bot, conn, message, user_data):
    user_data['equipment_release_date'] = message.text
    bot.send_message(message.chat.id, "Введите цвет оборудования:")
    bot.register_next_step_handler(message, lambda msg: get_color(bot, conn, msg, user_data))

def get_color(bot, conn, message, user_data):
    user_data['color'] = message.text
    bot.send_message(message.chat.id, "Введите цену оборудования:")
    bot.register_next_step_handler(message, lambda msg: get_price(bot, conn, msg, user_data))


def get_price(bot, conn, message, user_data):
    user_data['price'] = float(message.text)


    bot.send_message(message.chat.id, "Введите серийный номер оборудования:")
    bot.register_next_step_handler(message, lambda msg: get_serial_number(bot, conn, msg, user_data))


def get_serial_number(bot, conn, message, user_data):
    user_data['serial_number'] = message.text
    bot.send_message(message.chat.id, "Введите название файла для сохранения акта:")
    bot.register_next_step_handler(message, lambda msg: save_file(bot, conn, msg, user_data))


def save_file(bot, conn, message, user_data):
    user_data['file_name'] = message.text
    user_data['quantity'] = 1
    user_data['total_cost'] = user_data['price'] * user_data['quantity']


    user_data['serial_numbers'] = user_data['serial_number']

    file_path = create_asset_issue_act(user_data, user_data['file_name'])
    with open(file_path, 'rb') as document:
        bot.send_document(message.chat.id, document)

    bot.send_message(message.chat.id, "Акт выдачи актива успешно создан!")
